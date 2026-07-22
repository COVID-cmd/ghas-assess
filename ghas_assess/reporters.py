"""Output renderers: terminal, JSON, HTML, DOCX.

Each report presents, per check: STATUS, what is configured now (detail),
the best practice (what should be), and WHY it matters. Findings are grouped by category.
"""
import json
import datetime
import html as _html
import math
from .findings import FINDINGS, CATEGORIES, GAP, OK, INFO, UNKNOWN, SEVERITY_ORDER

SEV_COLOR = {"Critical": "C0392B", "High": "C0392B", "Medium": "BF8F00",
             "Low": "2E7D32", "Info": "5E7A90"}
STATUS_TXT = {GAP: "GAP", OK: "OK", INFO: "INFO", UNKNOWN: "MANUAL"}


def _cat(fid):
    return FINDINGS.get(fid, {}).get("category", "Other")


def _sort_key(r):
    return (CATEGORIES.index(_cat(r.finding_id)) if _cat(r.finding_id) in CATEGORIES else 99,
            0 if r.status == GAP else 1, SEVERITY_ORDER.get(r.severity, 9), r.finding_id, r.scope)


def _evidence_str(ev):
    """Compact human-readable rendering of the raw values that triggered a finding."""
    if not ev:
        return ""
    parts = []
    for k, v in ev.items():
        if isinstance(v, (list, tuple)):
            v = ", ".join(str(x) for x in v[:8]) + ("..." if len(v) > 8 else "")
        parts.append(f"{k} = {v}")
    return "; ".join(parts)


def _repo_scores(results):
    """Score each repo: count gaps by severity, compute a simple posture score (0-100)."""
    repos = {}
    for r in results:
        if not r.scope.startswith("repo:"):
            continue
        name = r.scope[len("repo:"):]
        d = repos.setdefault(name, {"High": 0, "Medium": 0, "Low": 0, "gaps": 0,
                                    "ok": 0, "checks": 0})
        d["checks"] += 1
        if r.status == GAP:
            d["gaps"] += 1
            if r.severity in d:
                d[r.severity] += 1
        elif r.status == OK:
            d["ok"] += 1
    # weighted penalty: High=5, Medium=2, Low=1; score = 100 - normalised penalty
    rows = []
    for name, d in repos.items():
        penalty = d["High"] * 5 + d["Medium"] * 2 + d["Low"] * 1
        score = max(0, 100 - penalty * 3)
        rows.append({"repo": name, "score": score, "penalty": penalty, **d})
    rows.sort(key=lambda x: (-x["penalty"], x["repo"]))  # worst first
    return rows


def _coverage_summary(results):
    """Break down the run: totals by status, and by automation confidence."""
    from .findings import AUTO, PARTIAL, MANUAL
    by_status = {}
    for r in results:
        by_status[r.status] = by_status.get(r.status, 0) + 1
    # unique checks exercised
    ids = sorted({r.finding_id for r in results})
    return {"total_results": len(results), "unique_checks": len(ids),
            "by_status": by_status, "ids": ids}


def _stacked_bars_svg(results, width=560):
    """Horizontal stacked bars: High/Medium/Low gaps per category."""
    cats = CATEGORIES
    data = {}
    for r in results:
        if r.status != GAP:
            continue
        c = FINDINGS.get(r.finding_id, {}).get("category")
        if not c:
            continue
        d = data.setdefault(c, {"High": 0, "Medium": 0, "Low": 0})
        s = r.severity if r.severity in d else "Low"
        d[s] += 1
    maxtot = max([sum(data.get(c, {}).values()) for c in cats] + [1])
    row_h, gap, label_w = 26, 14, 130
    x0 = label_w + 8
    barw = width - x0 - 40
    svg_h = len(cats) * (row_h + gap) + 20
    colors = {"High": "#e74c3c", "Medium": "#f39c12", "Low": "#2ecc71"}
    rows = []
    y = 10
    for c in cats:
        d = data.get(c, {"High": 0, "Medium": 0, "Low": 0})
        tot = sum(d.values())
        rows.append(f'<text x="{label_w}" y="{y+row_h/2+4}" fill="#cfe0ee" '
                    f'font-size="12" text-anchor="end" font-family="sans-serif">{c}</text>')
        x = x0
        for sev in ("High", "Medium", "Low"):
            w = (d[sev] / maxtot) * barw if maxtot else 0
            if w > 0:
                rows.append(f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{row_h}" '
                            f'fill="{colors[sev]}" rx="3"/>')
                if w > 16:
                    rows.append(f'<text x="{x+w/2:.1f}" y="{y+row_h/2+4}" fill="#0b1a26" '
                                f'font-size="11" font-weight="700" text-anchor="middle" '
                                f'font-family="sans-serif">{d[sev]}</text>')
                x += w + 2
        rows.append(f'<text x="{x0+barw+8}" y="{y+row_h/2+4}" fill="#9fb3c4" '
                    f'font-size="12" font-family="sans-serif">{tot}</text>')
        y += row_h + gap
    return (f'<svg viewBox="0 0 {width} {svg_h}" width="100%" height="{svg_h}" '
            f'xmlns="http://www.w3.org/2000/svg">{"".join(rows)}</svg>')


def _gauge_svg(score, size=200):
    """Circular posture-score gauge, 0-100, colour by band."""
    import math
    color = "2ecc71" if score >= 80 else "f39c12" if score >= 50 else "e74c3c"
    cx = cy = size / 2
    r = size * 0.4
    circ = 2 * math.pi * r
    frac = max(0, min(100, score)) / 100
    dash = circ * frac
    gap = circ - dash
    return (f'<svg viewBox="0 0 {size} {size}" width="{size}" height="{size}" '
            f'xmlns="http://www.w3.org/2000/svg">'
            f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="none" stroke="#16344c" '
            f'stroke-width="14"/>'
            f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="none" stroke="#{color}" '
            f'stroke-width="14" stroke-linecap="round" '
            f'stroke-dasharray="{dash:.1f} {gap:.1f}" '
            f'transform="rotate(-90 {cx} {cy})"/>'
            f'<text x="{cx}" y="{cy-4}" fill="#{color}" font-size="46" font-weight="800" '
            f'text-anchor="middle" font-family="sans-serif">{score}</text>'
            f'<text x="{cx}" y="{cy+22}" fill="#9fb3c4" font-size="13" '
            f'text-anchor="middle" font-family="sans-serif">posture score</text></svg>')


def _radar_svg(summary_by_cat, size=420):
    """Return an SVG radar/spiral chart of gap counts per category."""
    cats = CATEGORIES
    n = len(cats)
    cx = cy = size / 2
    R = size * 0.32
    maxv = max([summary_by_cat.get(c, 0) for c in cats] + [1])
    # rings
    rings = []
    for frac in (0.25, 0.5, 0.75, 1.0):
        pts = []
        for i in range(n):
            ang = -math.pi / 2 + 2 * math.pi * i / n
            x = cx + R * frac * math.cos(ang)
            y = cy + R * frac * math.sin(ang)
            pts.append(f"{x:.1f},{y:.1f}")
        rings.append(f'<polygon points="{" ".join(pts)}" fill="none" '
                     f'stroke="#1c3a52" stroke-width="1"/>')
    # spokes + labels
    spokes = []
    labels = []
    for i, c in enumerate(cats):
        ang = -math.pi / 2 + 2 * math.pi * i / n
        x = cx + R * math.cos(ang)
        y = cy + R * math.sin(ang)
        spokes.append(f'<line x1="{cx}" y1="{cy}" x2="{x:.1f}" y2="{y:.1f}" '
                      f'stroke="#1c3a52" stroke-width="1"/>')
        lx = cx + (R + 26) * math.cos(ang)
        ly = cy + (R + 26) * math.sin(ang)
        anchor = "middle"
        if lx < cx - 5:
            anchor = "end"
        elif lx > cx + 5:
            anchor = "start"
        val = summary_by_cat.get(c, 0)
        labels.append(f'<text x="{lx:.1f}" y="{ly:.1f}" fill="#cfe0ee" font-size="11" '
                      f'text-anchor="{anchor}" dominant-baseline="middle">{_html.escape(c)} '
                      f'({val})</text>')
    # data polygon
    dpts = []
    for i, c in enumerate(cats):
        ang = -math.pi / 2 + 2 * math.pi * i / n
        frac = summary_by_cat.get(c, 0) / maxv
        x = cx + R * frac * math.cos(ang)
        y = cy + R * frac * math.sin(ang)
        dpts.append(f"{x:.1f},{y:.1f}")
    data = (f'<polygon points="{" ".join(dpts)}" fill="rgba(192,57,43,0.35)" '
            f'stroke="#e74c3c" stroke-width="2"/>')
    dots = ""
    for p in dpts:
        x, y = p.split(",")
        dots += f'<circle cx="{x}" cy="{y}" r="3" fill="#e74c3c"/>'
    return (f'<svg viewBox="0 0 {size} {size}" width="{size}" height="{size}" '
            f'xmlns="http://www.w3.org/2000/svg">{"".join(rings)}{"".join(spokes)}'
            f'{data}{dots}{"".join(labels)}'
            f'<text x="{cx}" y="16" fill="#02fff0" font-size="13" text-anchor="middle" '
            f'font-weight="700">Gaps by category</text></svg>')


# ---------------- terminal ----------------
def to_terminal(results, summary):
    C = {"red": "\033[91m", "yel": "\033[93m", "grn": "\033[92m", "gry": "\033[90m",
         "bold": "\033[1m", "end": "\033[0m", "cyan": "\033[96m"}
    L = []
    L.append(f"\n{C['bold']}{C['cyan']}GHAS Security Assessment - {summary['org']}{C['end']}")
    L.append(f"{C['gry']}{summary['repos_assessed']} repos - {summary['gaps']} gaps "
             f"of {summary['total_results']} checks{C['end']}")
    bits = []
    for sev, n in summary["gaps_by_severity"].items():
        col = C["red"] if sev in ("Critical", "High") else C["yel"] if sev == "Medium" else C["grn"]
        bits.append(f"{col}{sev}: {n}{C['end']}")
    if bits:
        L.append("  " + "  ".join(bits))

    last_cat = None
    for r in sorted(results, key=_sort_key):
        cat = _cat(r.finding_id)
        if cat != last_cat:
            L.append(f"\n{C['bold']}{cat}{C['end']}")
            L.append(C["gry"] + "-" * 96 + C["end"])
            last_cat = cat
        if r.status == GAP:
            col = C["red"] if r.severity in ("Critical", "High") else C["yel"] if r.severity == "Medium" else C["grn"]
        elif r.status == OK:
            col = C["grn"]
        else:
            col = C["gry"]
        scope = (r.scope[:18] + "..") if len(r.scope) > 20 else r.scope
        st = STATUS_TXT[r.status]
        detail = (r.detail[:50] + "...") if len(r.detail) > 53 else r.detail
        L.append(f"{col}{r.finding_id:<24}{st:<7}{r.severity:<7}{C['end']}{scope:<21} {detail}")
    return "\n".join(L)


# ---------------- json ----------------
def to_json(results, summary, path):
    payload = {
        "generated": datetime.datetime.now(datetime.timezone.utc).isoformat(),
        "summary": summary,
        "results": [],
    }
    for r in sorted(results, key=_sort_key):
        d = r.to_dict()
        meta = FINDINGS.get(r.finding_id, {})
        d["category"] = meta.get("category")
        d["best_practice"] = meta.get("best_practice")
        d["why"] = meta.get("why")
        d["improvement"] = meta.get("improvement") if r.status == GAP else None
        payload["results"].append(d)
    with open(path, "w") as f:
        json.dump(payload, f, indent=2)
    return path


# ---------------- html ----------------
def to_html(results, summary, path):
    gen = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    badge_bg = {"GAP": "#c0392b", "OK": "#2e7d32", "INFO": "#456", "MANUAL": "#7d5ba6"}

    # gaps per category for the radar chart
    gaps_by_cat = {}
    configured = []  # OK results
    for r in results:
        if r.status == GAP:
            gaps_by_cat[_cat(r.finding_id)] = gaps_by_cat.get(_cat(r.finding_id), 0) + 1
        elif r.status == OK:
            configured.append(r)
    radar = _radar_svg(gaps_by_cat)

    # "what is already configured" section
    conf_rows = []
    for r in sorted(configured, key=_sort_key):
        meta = FINDINGS.get(r.finding_id, {})
        conf_rows.append(f"""
        <tr><td class="fid">{r.finding_id}</td>
        <td class="scope">{_html.escape(r.scope)}</td>
        <td><b>{_html.escape(meta.get('title',''))}</b> &mdash; {_html.escape(r.detail)}</td></tr>""")
    configured_block = ""
    if conf_rows:
        configured_block = f"""
        <h2 style="color:#2ecc71">What is already configured <span class="cntr">{len(conf_rows)} OK</span></h2>
        <table><thead><tr><th>ID</th><th>Scope</th><th>Configured correctly</th></tr></thead>
        <tbody>{''.join(conf_rows)}</tbody></table>"""

    # gap/other sections by category
    sections = []
    for cat in CATEGORIES:
        rows = [r for r in sorted(results, key=_sort_key)
                if _cat(r.finding_id) == cat and r.status != OK]
        if not rows:
            continue
        gap_n = sum(1 for r in rows if r.status == GAP)
        body = []
        for r in rows:
            meta = FINDINGS.get(r.finding_id, {})
            badge = STATUS_TXT[r.status]
            color = SEV_COLOR.get(r.severity, "5E7A90")
            why = meta.get("why", "")
            best = meta.get("best_practice", "")
            imp = meta.get("improvement", "") if r.status == GAP else ""
            action = (f'<div class="imp"><b>What to do:</b> {_html.escape(imp)}</div>' if imp else "")
            ev = _evidence_str(r.evidence)
            eb = (f'<div class="ev"><b>Why flagged (evidence):</b> <code>{_html.escape(ev)}</code></div>'
                  if ev and r.status == GAP else "")
            body.append(f"""
            <tr>
              <td class="fid">{r.finding_id}</td>
              <td><span class="badge" style="background:{badge_bg[badge]}">{badge}</span></td>
              <td style="color:#{color};font-weight:600">{_html.escape(r.severity)}</td>
              <td class="scope">{_html.escape(r.scope)}</td>
              <td>
                <div class="tt">{_html.escape(meta.get('title',''))}</div>
                <div class="now"><b>Now:</b> {_html.escape(r.detail)}</div>
                <div class="bp"><b>Best practice:</b> {_html.escape(best)}</div>
                <div class="why"><b>Why it matters:</b> {_html.escape(why)}</div>
                {eb}{action}
              </td>
            </tr>""")
        sections.append(f"""
        <h2>{_html.escape(cat)} <span class="cntr">{gap_n} gap(s)</span></h2>
        <table><thead><tr><th>ID</th><th>Status</th><th>Sev</th><th>Scope</th>
        <th>What is configured, what should be, why, and the evidence</th></tr></thead>
        <tbody>{''.join(body)}</tbody></table>""")

    chips = "".join(
        f'<span class="chip" style="background:#{SEV_COLOR.get(s,"555")}">{s}: {n}</span>'
        for s, n in summary["gaps_by_severity"].items())

    # dashboard pieces
    repo_rows = _repo_scores(results)
    total_checks = summary["total_results"]
    ok_count = sum(1 for r in results if r.status == OK)
    gap_count = summary["gaps"]
    manual_count = sum(1 for r in results if r.status == UNKNOWN)
    high_count = summary["gaps_by_severity"].get("High", 0) + summary["gaps_by_severity"].get("Critical", 0)

    # overall org posture score: average of repo scores (fallback to gap-weighted)
    if repo_rows:
        org_score = round(sum(r["score"] for r in repo_rows) / len(repo_rows))
    else:
        org_score = max(0, 100 - gap_count * 2)
    gauge = _gauge_svg(org_score)
    band = ("Strong" if org_score >= 80 else "Needs work" if org_score >= 50 else "At risk")
    band_color = "2ecc71" if org_score >= 80 else "f39c12" if org_score >= 50 else "e74c3c"

    def score_color(s):
        return "2ecc71" if s >= 80 else "f39c12" if s >= 50 else "e74c3c"

    repo_table = "".join(
        f"""<tr>
          <td class="scope">{_html.escape(rw['repo'])}</td>
          <td><span class="scorepill" style="background:#{score_color(rw['score'])}1a;color:#{score_color(rw['score'])}">{rw['score']}</span></td>
          <td style="color:#e74c3c;font-weight:500">{rw['High']}</td>
          <td style="color:#f39c12;font-weight:500">{rw['Medium']}</td>
          <td style="color:#2ecc71;font-weight:500">{rw['Low']}</td>
          <td style="color:#2ecc71">{rw['ok']}</td>
        </tr>""" for rw in repo_rows)

    # category score bars
    cat_bars = []
    for cat in CATEGORIES:
        g = gaps_by_cat.get(cat, 0)
        catchecks = sum(1 for r in results if _cat(r.finding_id) == cat)
        pct = int(100 * g / catchecks) if catchecks else 0
        barcol = "e74c3c" if pct >= 50 else "f39c12" if pct >= 20 else "2ecc71"
        cat_bars.append(f"""
        <div class="catbar">
          <div class="catname">{_html.escape(cat)}</div>
          <div class="bartrack"><div class="barfill" style="width:{max(pct,2)}%;background:#{barcol}"></div></div>
          <div class="catval">{g}</div>
        </div>""")

    # priority remediation plan: group gaps by finding, count AFFECTED REPOS correctly
    gap_results = [r for r in results if r.status == GAP]
    by_finding = {}
    for r in gap_results:
        fid = r.finding_id
        d = by_finding.setdefault(fid, {"severity": r.severity, "scopes": set(),
                                        "org": False})
        if r.scope.startswith("repo:"):
            d["scopes"].add(r.scope[len("repo:"):])
        else:
            d["org"] = True
        # keep worst severity seen
        if SEVERITY_ORDER.get(r.severity, 9) < SEVERITY_ORDER.get(d["severity"], 9):
            d["severity"] = r.severity

    def _plan_sort(item):
        fid, d = item
        affected = len(d["scopes"]) + (1 if d["org"] else 0)
        return (SEVERITY_ORDER.get(d["severity"], 9), -affected)

    plan = sorted(by_finding.items(), key=_plan_sort)
    total_repos = summary["repos_assessed"] or 1

    plan_rows = []
    for fid, d in plan:
        meta = FINDINGS.get(fid, {})
        affected = len(d["scopes"]) if not d["org"] else total_repos
        # scope of fix: org-level setting fixes everything at once = a "quick win"
        is_orgfix = meta.get("scope") == "org" or d["org"]
        fixtype = ('<span class="tag tagwin">org-wide fix</span>' if is_orgfix
                   else '<span class="tag tagrepo">per-repo</span>')
        sev = d["severity"]
        sevcol = SEV_COLOR.get(sev, "5E7A90")
        # coverage bar: how much of the estate is affected
        pct = int(100 * affected / total_repos) if not d["org"] else 100
        plan_rows.append(f"""
        <tr>
          <td><span class="sevdot" style="background:#{sevcol}"></span><span style="color:#{sevcol};font-weight:600">{sev}</span></td>
          <td><div class="pfind">{_html.escape(meta.get('title', fid))}</div><div class="pfid">{_html.escape(fid)}</div></td>
          <td><div class="affbar"><div class="afffill" style="width:{max(pct,6)}%;background:#{sevcol}"></div></div>
              <div class="affn">{affected}/{total_repos} repos</div></td>
          <td>{fixtype}</td>
          <td>{_html.escape(meta.get('improvement', ''))}</td>
        </tr>""")
    plan_rows = "".join(plan_rows[:10])

    # quick wins: org-level fixes that clear gaps across many repos at once
    quick = [(fid, d) for fid, d in plan
             if (FINDINGS.get(fid, {}).get("scope") == "org" or d["org"])]
    quick_count = len(quick)

    stacked = _stacked_bars_svg(results)
    dashboard = f"""
    <div class="herorow">
      <div class="hero">
        <div class="gauge">{gauge}</div>
        <div class="heroband" style="color:#{band_color}">{band}</div>
      </div>
      <div class="herostats">
        <div class="statgrid">
          <div class="stat"><div class="snum" style="color:#5dd6ff">{summary['repos_assessed']}</div><div class="slab">Repositories</div></div>
          <div class="stat"><div class="snum" style="color:#5dd6ff">{total_checks}</div><div class="slab">Checks run</div></div>
          <div class="stat statbad"><div class="snum" style="color:#e74c3c">{gap_count}</div><div class="slab">Gaps found</div></div>
          <div class="stat"><div class="snum" style="color:#e74c3c">{high_count}</div><div class="slab">High severity</div></div>
          <div class="stat statgood"><div class="snum" style="color:#2ecc71">{ok_count}</div><div class="slab">Configured OK</div></div>
          <div class="stat statwin"><div class="snum" style="color:#02fff0">{quick_count}</div><div class="slab">Org-wide quick wins</div></div>
        </div>
        <div class="chips">{chips}</div>
      </div>
    </div>
    <div class="dashgrid">
      <div class="panelbox">
        <h3>Gap severity by category</h3>
        <div class="legend"><span><i class="sw" style="background:#e74c3c"></i>High</span><span><i class="sw" style="background:#f39c12"></i>Medium</span><span><i class="sw" style="background:#2ecc71"></i>Low</span></div>
        {stacked}
      </div>
      <div class="panelbox">
        <h3>Where gaps concentrate</h3>
        <div class="radar">{radar}</div>
      </div>
    </div>
    <div class="panelbox">
      <h3>Priority remediation plan <span class="hint">severity first, then blast radius</span></h3>
      <p class="planhint">Start at the top. <b style="color:#02fff0">Org-wide fixes</b> clear the gap across every affected repository in a single change &mdash; the fastest way to raise the score.</p>
      <table class="plan"><thead><tr><th>Severity</th><th>Finding</th><th>Estate impact</th><th>Fix type</th><th>Action</th></tr></thead>
      <tbody>{plan_rows}</tbody></table>
    </div>
    <div class="panelbox">
      <h3>Repositories ranked by posture <span class="hint">worst first</span></h3>
      <table><thead><tr><th>Repository</th><th>Score</th><th>High</th><th>Medium</th><th>Low</th><th>OK</th></tr></thead>
      <tbody>{repo_table}</tbody></table>
    </div>
    """

    details = f"{_status_legend_html()}{configured_block}{''.join(sections)}{_checks_performed_html(results)}"

    doc = f"""<!doctype html><html><head><meta charset="utf-8">
<title>GHAS Security Assessment - {_html.escape(str(summary['org']))}</title>
<style>
 body{{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#0b1a26;color:#e8ecf1;margin:0;padding:0;line-height:1.45}}
 .wrap{{padding:28px 32px}}
 h1{{color:#02fff0;margin:0 0 4px}} h2{{color:#b997ff;margin:26px 0 8px;border-bottom:1px solid #16344c;padding-bottom:4px}}
 h3{{color:#9fe0ff;margin:20px 0 8px;font-size:15px}}
 .cntr{{font-size:13px;color:#9fb3c4;font-weight:400}}
 .sub{{color:#9fb3c4;margin-bottom:12px}}
 .tabs{{display:flex;gap:6px;background:#0e2333;padding:10px 32px 0;position:sticky;top:0;z-index:5;border-bottom:1px solid #16344c}}
 .tab{{padding:10px 20px;cursor:pointer;border-radius:8px 8px 0 0;color:#9fb3c4;font-weight:600}}
 .tab.active{{background:#0b1a26;color:#02fff0}}
 .panel{{display:none}} .panel.active{{display:block}}
 .cards{{display:flex;gap:14px;flex-wrap:wrap;margin:8px 0 16px}}
 .herorow{{display:flex;gap:20px;flex-wrap:wrap;align-items:stretch;margin:12px 0 18px}}
 .hero{{background:#0e2333;border-radius:16px;padding:20px 28px;display:flex;flex-direction:column;align-items:center;justify-content:center;min-width:220px}}
 .heroband{{font-size:15px;font-weight:700;margin-top:6px;letter-spacing:.3px}}
 .herostats{{flex:1;min-width:320px;display:flex;flex-direction:column;justify-content:center}}
 .statgrid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:12px}}
 .stat{{background:#0e2333;border-radius:12px;padding:14px 16px;text-align:center;border-top:3px solid #16344c}}
 .stat.statbad{{border-top-color:#e74c3c}} .stat.statgood{{border-top-color:#2ecc71}} .stat.statwin{{border-top-color:#02fff0}}
 .legend{{display:flex;gap:16px;margin:0 0 10px;font-size:12px;color:#9fb3c4}}
 .legend i.sw{{display:inline-block;width:11px;height:11px;border-radius:3px;margin-right:5px;vertical-align:-1px}}
 .planhint{{color:#9fb3c4;font-size:12.5px;margin:0 0 12px}}
 .plan td{{vertical-align:middle}}
 .sevdot{{display:inline-block;width:9px;height:9px;border-radius:50%;margin-right:6px;vertical-align:0}}
 .pfind{{font-weight:600;color:#e8ecf1}} .pfid{{font-size:11px;color:#5e7a90;font-family:ui-monospace,monospace}}
 .affbar{{width:100px;height:8px;background:#081826;border-radius:4px;overflow:hidden;margin-bottom:3px}}
 .afffill{{height:100%;border-radius:4px}} .affn{{font-size:11px;color:#9fb3c4}}
 .tag{{font-size:11px;padding:3px 9px;border-radius:10px;font-weight:700;white-space:nowrap}}
 .tagwin{{background:#02fff01a;color:#02fff0}} .tagrepo{{background:#9fb3c41a;color:#9fb3c4}}
 .snum{{font-size:30px;font-weight:800;line-height:1}} .slab{{color:#9fb3c4;font-size:11.5px;margin-top:5px}}
 .chips{{margin:14px 0 4px}} .chip{{display:inline-block;color:#fff;padding:5px 12px;border-radius:14px;margin-right:8px;font-size:13px;font-weight:700}}
 .dashgrid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(320px,1fr));gap:16px;margin-bottom:16px}}
 .panelbox{{background:#0e2333;border-radius:16px;padding:16px 20px}}
 .panelbox h3{{margin:2px 0 12px}} .hint{{font-size:12px;color:#5e7a90;font-weight:400}}
 .radar{{display:flex;justify-content:center}}
 .catbar{{display:flex;align-items:center;gap:12px;margin:11px 0}}
 .catname{{width:140px;font-size:13px;color:#cfe0ee}}
 .bartrack{{flex:1;height:14px;background:#081826;border-radius:7px;overflow:hidden}}
 .barfill{{height:100%;border-radius:7px;transition:width .3s}}
 .catval{{width:28px;text-align:right;font-size:14px;color:#e8ecf1;font-weight:700}}
 .scorepill{{padding:3px 12px;border-radius:12px;font-weight:800;font-size:14px}}
 .afx{{background:#e74c3c22;color:#e74c3c;padding:2px 9px;border-radius:10px;font-weight:700;font-size:12px}}
 .panelbox table{{margin-bottom:0;background:transparent}}
 .panelbox th{{background:#0b2033}}
 table{{border-collapse:collapse;width:100%;background:#0e2333;border-radius:10px;overflow:hidden;margin-bottom:8px}}
 th{{background:#12304a;text-align:left;padding:9px 12px;font-size:12.5px;color:#cfe0ee}}
 td{{padding:9px 12px;border-top:1px solid #16344c;font-size:13px;vertical-align:top}}
 .fid{{font-weight:700;color:#02fff0}} .scope{{color:#9fb3c4;font-family:ui-monospace,monospace;font-size:11.5px}}
 .badge{{color:#fff;padding:2px 8px;border-radius:6px;font-size:11px;font-weight:700}}
 .tt{{font-weight:700;margin-bottom:3px}}
 .now{{margin:2px 0}} .bp{{color:#9fe0c4;margin:2px 0}} .why{{color:#c9b8ff;margin:2px 0}}
 .ev{{color:#9fb3c4;margin:3px 0}} .ev code{{background:#0b1a26;padding:1px 5px;border-radius:4px;font-size:11.5px}}
 .imp{{color:#ffd479;margin-top:4px}}
 .catrow{{background:#0b2436;color:#b997ff;font-weight:700;font-size:12px}}
 .legendbox{{background:#0e2333;border-radius:16px;padding:16px 20px;margin-bottom:18px}}
 .legendbox h3{{margin:2px 0 12px}}
 .lgrow{{display:flex;align-items:flex-start;gap:12px;margin:8px 0}}
 .lgrow .badge{{min-width:64px;text-align:center;flex-shrink:0}}
 .lgtext{{font-size:13px;color:#cfe0ee}}
 .lgsev{{margin-top:12px;padding-top:12px;border-top:1px solid #16344c;font-size:13px;color:#cfe0ee}}
 .foot{{color:#5e7a90;margin-top:22px;font-size:12px}}
</style></head><body>
<div class="tabs">
  <div class="tab active" onclick="showTab('dash')">Dashboard</div>
  <div class="tab" onclick="showTab('detail')">Detailed findings</div>
</div>
<div class="wrap">
<h1>GitHub Advanced Security - Assessment</h1>
<div class="sub">Organisation <b>{_html.escape(str(summary['org']))}</b> &middot; {summary['repos_assessed']} repositories &middot; {summary['gaps']} gaps of {summary['total_results']} checks &middot; generated {gen}</div>
<div id="dash" class="panel active">{dashboard}</div>
<div id="detail" class="panel">{details}</div>
<div class="foot">Generated {gen} by ghas-assess &middot; read-only &middot; MANUAL = not exposed by the API, verify in the UI.</div>
</div>
<script>
function showTab(id){{
  document.querySelectorAll('.panel').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.tab').forEach(t=>t.classList.remove('active'));
  document.getElementById(id).classList.add('active');
  event.target.classList.add('active');
}}
</script>
</body></html>"""
    with open(path, "w") as f:
        f.write(doc)
    return path


def _status_legend_html():
    """Explain the four statuses and severities so any reader understands the report."""
    items = [
        ("GAP", "#c0392b", "Something is misconfigured and needs fixing. These are the findings."),
        ("OK", "#2e7d32", "Checked and correctly configured. A win \u2014 no action needed."),
        ("MANUAL", "#7d5ba6", "The GitHub API cannot read this setting; verify it yourself in the UI "
                              "(or supply a manual-answers file)."),
        ("INFO", "#456", "No problem, just context worth knowing (e.g. open alerts that aren't "
                         "critical). Not an action item."),
    ]
    rows = "".join(
        f'<div class="lgrow"><span class="badge" style="background:{c}">{name}</span>'
        f'<span class="lgtext">{_html.escape(desc)}</span></div>' for name, c, desc in items)
    sev = ('<div class="lgsev"><b>Severity</b> (on gaps): '
           '<span style="color:#e74c3c">High</span> \u00b7 '
           '<span style="color:#f39c12">Medium</span> \u00b7 '
           '<span style="color:#2ecc71">Low</span> \u2014 how urgent the fix is.</div>')
    return f"""
    <div class="legendbox">
      <h3>How to read this report</h3>
      {rows}{sev}
    </div>"""


def _checks_performed_html(results):
    """A catalogue section: every check the tool ran, grouped by category, with how many
    times it was exercised and how its state is determined."""
    cov = _coverage_summary(results)
    # count per finding id
    per_id = {}
    for r in results:
        per_id.setdefault(r.finding_id, {"OK": 0, "GAP": 0, "INFO": 0, "UNKNOWN": 0})
        per_id[r.finding_id][r.status] = per_id[r.finding_id].get(r.status, 0) + 1
    rows = []
    for cat in CATEGORIES:
        cat_ids = [fid for fid in FINDINGS if FINDINGS[fid]["category"] == cat and fid in per_id]
        if not cat_ids:
            continue
        rows.append(f'<tr><td colspan="5" class="catrow">{_html.escape(cat)}</td></tr>')
        for fid in sorted(cat_ids):
            m = FINDINGS[fid]
            c = per_id[fid]
            tally = ", ".join(f"{k} {v}" for k, v in c.items() if v)
            rows.append(f"""<tr>
              <td class="fid">{fid}</td>
              <td>{_html.escape(m['title'])}</td>
              <td class="scope">{m['scope']}</td>
              <td>{m['automation']}</td>
              <td>{tally}</td></tr>""")
    sb = cov["by_status"]
    status_line = " &middot; ".join(
        f"{STATUS_TXT.get(k,k)}: {v}" for k, v in
        sorted(sb.items(), key=lambda kv: kv[0]))
    return f"""
    <h2>Checks performed <span class="cntr">{cov['unique_checks']} distinct checks &middot; {cov['total_results']} evaluations</span></h2>
    <div class="sub">Each repository and the organisation are evaluated against the catalogue below.
    {cov['total_results']} total evaluations this run &mdash; {status_line}.
    <b>How determined:</b> AUTO = read from the API &middot; PARTIAL = API signal, confirm nuance &middot; MANUAL = not exposed by the API.</div>
    <table><thead><tr><th>ID</th><th>Check</th><th>Scope</th><th>How</th><th>Results this run</th></tr></thead>
    <tbody>{''.join(rows)}</tbody></table>"""


# ---------------- docx ----------------
def to_docx(results, summary, path):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    t = doc.add_paragraph()
    run = t.add_run("GitHub Advanced Security - Security Best-Practice Assessment")
    run.bold = True; run.font.size = Pt(19); run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    s = doc.add_paragraph(
        f"Organisation {summary['org']}  -  {summary['repos_assessed']} repositories  -  "
        f"{summary['gaps']} gaps of {summary['total_results']} checks")
    s.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    gen = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Generated {gen} - read-only assessment").runs[0].italic = True
    if summary["gaps_by_severity"]:
        p = doc.add_paragraph(); p.add_run("Gaps by severity: ").bold = True
        p.add_run(", ".join(f"{k} {v}" for k, v in summary["gaps_by_severity"].items()))

    # how to read this report - status legend
    doc.add_heading("How to read this report", level=1)
    legend = [
        ("GAP", "Something is misconfigured and needs fixing. These are the findings."),
        ("OK", "Checked and correctly configured. A win - no action needed."),
        ("MANUAL", "The GitHub API cannot read this setting; verify it yourself in the UI "
                   "or supply a manual-answers file."),
        ("INFO", "No problem, just context worth knowing (e.g. open alerts that aren't "
                 "critical). Not an action item."),
    ]
    for name, desc in legend:
        pl = doc.add_paragraph(style="List Bullet")
        pl.add_run(f"{name} \u2014 ").bold = True
        pl.add_run(desc)
    ps = doc.add_paragraph()
    ps.add_run("Severity (on gaps): ").bold = True
    ps.add_run("High, Medium, Low - how urgent the fix is.")

    # what is already configured
    configured = [r for r in sorted(results, key=_sort_key) if r.status == OK]
    if configured:
        doc.add_heading("What is already configured", level=1)
        ct = doc.add_table(rows=1, cols=3); ct.style = "Light Grid Accent 1"
        for i, h in enumerate(["ID", "Scope", "Configured correctly"]):
            ct.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for r in configured:
            meta = FINDINGS.get(r.finding_id, {})
            c = ct.add_row().cells
            c[0].text = r.finding_id
            c[1].text = r.scope
            c[2].paragraphs[0].add_run(meta.get("title", "") + " - ").bold = True
            c[2].paragraphs[0].add_run(r.detail)

    for cat in CATEGORIES:
        rows = [r for r in sorted(results, key=_sort_key)
                if _cat(r.finding_id) == cat and r.status != OK]
        if not rows:
            continue
        doc.add_heading(cat, level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["ID", "Status", "Scope", "What is / should be / why / evidence"]):
            tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for r in rows:
            meta = FINDINGS.get(r.finding_id, {})
            c = tbl.add_row().cells
            c[0].text = r.finding_id
            c[1].text = f"{STATUS_TXT[r.status]} ({r.severity})"
            c[2].text = r.scope
            cell = c[3]
            cell.paragraphs[0].add_run(meta.get("title", "") + "\n").bold = True
            cell.add_paragraph().add_run("Now: ").bold = True
            cell.paragraphs[-1].add_run(r.detail)
            bp = cell.add_paragraph(); bp.add_run("Best practice: ").bold = True
            bp.add_run(meta.get("best_practice", ""))
            wy = cell.add_paragraph(); wy.add_run("Why: ").bold = True
            wr = wy.add_run(meta.get("why", "")); wr.italic = True
            ev = _evidence_str(r.evidence)
            if ev and r.status == GAP:
                ep = cell.add_paragraph(); ep.add_run("Why flagged (evidence): ").bold = True
                er = ep.add_run(ev); er.font.name = "Consolas"
            if r.status == GAP:
                ip = cell.add_paragraph(); ip.add_run("What to do: ").bold = True
                ir = ip.add_run(meta.get("improvement", ""))
                ir.font.color.rgb = RGBColor(0x2E, 0x5F, 0xA3)

    # checks performed catalogue
    cov = _coverage_summary(results)
    per_id = {}
    for r in results:
        per_id.setdefault(r.finding_id, {})
        per_id[r.finding_id][r.status] = per_id[r.finding_id].get(r.status, 0) + 1
    doc.add_heading("Checks performed", level=1)
    sb = cov["by_status"]
    status_line = ", ".join(f"{STATUS_TXT.get(k,k)}: {v}" for k, v in sorted(sb.items()))
    intro = doc.add_paragraph()
    intro.add_run(f"{cov['unique_checks']} distinct checks, {cov['total_results']} total "
                  f"evaluations this run ({status_line}). ")
    intro.add_run("How determined: AUTO = read from the API; PARTIAL = API signal, confirm "
                  "nuance; MANUAL = not exposed by the API.").italic = True
    ct = doc.add_table(rows=1, cols=5); ct.style = "Light Grid Accent 1"
    for i, h in enumerate(["ID", "Check", "Scope", "How", "Results this run"]):
        ct.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for cat in CATEGORIES:
        cat_ids = sorted(fid for fid in FINDINGS
                         if FINDINGS[fid]["category"] == cat and fid in per_id)
        if not cat_ids:
            continue
        hdr = ct.add_row().cells
        hr = hdr[0].paragraphs[0].add_run(cat); hr.bold = True
        for fid in cat_ids:
            m = FINDINGS[fid]; c = per_id[fid]
            tally = ", ".join(f"{STATUS_TXT.get(k,k)} {v}" for k, v in c.items() if v)
            row = ct.add_row().cells
            row[0].text = fid
            row[1].text = m["title"]
            row[2].text = m["scope"]
            row[3].text = m["automation"]
            row[4].text = tally

    note = doc.add_paragraph(
        "MANUAL = state not exposed by the GitHub API; verify in the UI or supply a manual-answers "
        "file. This tool is strictly read-only.")
    note.runs[0].italic = True; note.runs[0].font.size = Pt(9)
    doc.save(path)
    return path
